import fitz
import docx
import io
import logging

# OCR fallback is optional and requires external binaries
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

from app.utils.helpers import setup_logger

logger = setup_logger("cv_parser")

def _ocr_fallback(file_bytes: bytes, text: str) -> str:
    """Helper to attempt OCR fallback."""
    if not HAS_OCR:
        logger.warning("No text extracted and OCR dependencies are missing.")
        return text
    logger.info("No text extracted using PyMuPDF. Attempting OCR fallback...")
    try:
        images = convert_from_bytes(file_bytes)
        for image in images:
            text += pytesseract.image_to_string(image)
    except Exception as ocr_e:
        logger.error(f"OCR fallback failed: {ocr_e}")
    return text

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file.

    Handles encrypted PDFs, resets the stream, and enforces a size limit (5 MB).
    Falls back to OCR if no text is found and OCR dependencies are installed.
    """
    MAX_SIZE = 5 * 1024 * 1024  # 5 MB
    if len(file_bytes) > MAX_SIZE:
        logger.error("PDF size exceeds 5 MB limit.")
        raise ValueError("PDF file is too large; limit is 5 MB.")
    text = ""
    try:
        # Use BytesIO to ensure the stream is at the start
        doc = fitz.open(stream=io.BytesIO(file_bytes), filetype="pdf")
        if doc.is_encrypted:
            try:
                doc.authenticate("")
            except Exception as auth_e:
                logger.error(f"Failed to decrypt PDF: {auth_e}")
                raise ValueError("Encrypted PDF cannot be processed.")
        for page in doc:
            text += page.get_text()
        # If no text extracted, attempt OCR fallback
        if not text.strip():
            return _ocr_fallback(file_bytes, text)
    except Exception as e:
        logger.exception("Error extracting text from PDF")
        raise
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    text = ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Extract text from tables to prevent losing table-based CV content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text += para.text + " "
                    text += "\n"
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
    return text

def parse_cv(file_bytes: bytes, filename: str) -> str:
    logger.info(f"Parsing CV: {filename}")
    if filename.lower().endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename.lower().endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        logger.warning(f"Unsupported file format: {filename}")
        return ""
